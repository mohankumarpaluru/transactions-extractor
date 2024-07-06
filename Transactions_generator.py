import re
import tkinter as tk
from datetime import datetime
from pathlib import Path
from tkinter import filedialog, messagebox, ttk

import pandas as pd
import pytesseract
import sv_ttk
from docx import Document
from docx.shared import Inches
from PIL import Image

# Path to tesseract executable (change this if necessary)
pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"

DEFAULT_INPUT_DIR = str(Path(".").joinpath("Screenshots").absolute())
DEFAULT_OUTPUT_FILE = str(Path(".").joinpath("Transaction.xlsx").absolute())


def get_image_files(directory):
    image_extensions = {".jpeg", ".jpg", ".png"}
    image_files = [
        file for file in Path(directory).rglob("*") if file.suffix.lower() in image_extensions
    ]
    return image_files

def create_docx_file(df, output_file):
    doc = Document()
    doc.styles['Normal'].font.name = 'Aptos'

    # Step 3: Add a title to the document
    doc.add_heading('something here', level=1)
    # Step 4: Add the DataFrame as a table
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = 'Table Grid'
    # Step 5: Add the header row
    hdr_cells = table.rows[0].cells
    for i, column in enumerate(df.columns):
        hdr_cells[i].text = column

    # Step 6: Add the DataFrame data to the table
    for index, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, cell in enumerate(row):
            row_cells[i].text = str(cell)

    # Step 7: Save the document
    doc.save(output_file.replace(".xlsx", '.docx'))

def process_image(image_path):
    # Open the image file
    img = Image.open(str(image_path))
    # Use pytesseract to do OCR on the image
    text = pytesseract.image_to_string(img)
    return text


def append_transaction_row(text, df):
    # Define regular expressions to extract the required information
    time_pattern = re.compile(r"(\d{1,2}:\d{2} [APM]{2})")
    date_pattern = re.compile(r"(\d{1,2} [A-Za-z]{3} \d{4})")
    transaction_id_pattern = re.compile(r"(T\d{22})")
    utr_pattern = re.compile(r"UTR: (\d{12})")
    amount_pattern = re.compile(r"=\s*([\d,]+)")

    # Search for the patterns in the text
    time_match = time_pattern.search(text)
    date_match = date_pattern.search(text)
    transaction_id_match = transaction_id_pattern.search(text)
    utr_match = utr_pattern.search(text)
    amount_match = amount_pattern.search(text)

    # Extract the matched values
    time = time_match.group(1) if time_match else None
    date = date_match.group(1) if date_match else None
    transaction_id = transaction_id_match.group(1) if transaction_id_match else None
    utr = utr_match.group(1) if utr_match else None
    amount = amount_match.group(1).replace(",", "") if amount_match else None

    # Convert the date to the desired format
    if date:
        date_obj = datetime.strptime(date, "%d %b %Y")
        formatted_date = date_obj.strftime("%d %b %Y")
    else:
        formatted_date = None

    dt_str = f"{formatted_date} {time}"
    dt_obj = datetime.strptime(dt_str, "%d %b %Y %I:%M %p")
    epoch_time =  datetime.timestamp(dt_obj)

    # Create a DataFrame for the new data
    new_data = {
        "Date": [formatted_date],
        "Time": [time],
        "Transaction ID": [transaction_id],
        "UTR": [utr],
        "Amount": [amount],
        "EpochTime": [epoch_time]
    }
    new_df = pd.DataFrame(new_data)

    # Append the new data to the existing DataFrame
    updated_df = pd.concat([df, new_df], ignore_index=True)

    return updated_df


def process_transactions(input_dir, output_file):
    transaction_df = pd.DataFrame(columns=["Date", "Time", "Transaction ID", "UTR", "Amount", "EpochTime"])
    images = get_image_files(input_dir)

    for image in images:
        text = process_image(image)
        transaction_df = append_transaction_row(text, transaction_df)

    transaction_df = transaction_df.sort_values(by="EpochTime")
    transaction_df = transaction_df.drop(columns=["EpochTime"])

    transaction_df["Amount"] = pd.to_numeric(transaction_df["Amount"])

    # Sum the "Amount" column
    total_amount = transaction_df["Amount"].sum()

    # Create a new row with the sum of the "Amount" column
    new_row = {
        "Date": [""],
        "Time": [""],
        "Transaction ID":[""],
        "UTR": ["Total"],
        "Amount": [total_amount]
    }

    df_total = pd.DataFrame(new_row)
    # Append the new row to the DataFrame
    transaction_df = pd.concat([transaction_df, df_total], ignore_index=True)

    transaction_df.to_excel(output_file, index=False)
    create_docx_file(transaction_df, output_file)
    messagebox.showinfo("Success", f"Transactihttps://chatgpt.com/c/948f65dc-e91a-4900-af25-90a6c67fbe55ons have been processed and saved to {output_file}")


def select_input_directory():
    directory = filedialog.askdirectory(initialdir=DEFAULT_INPUT_DIR)
    if directory:
        input_dir_var.set(directory)


def select_output_file():
    file = filedialog.asksaveasfilename(defaultextension=".xlsx",
                                        filetypes=[("Excel files", "*.xlsx"), ("All files", "*.*")],
                                        initialfile=DEFAULT_OUTPUT_FILE)
    if file:
        output_file_var.set(file)


def create_gui():
    global input_dir_var, output_file_var
    root = tk.Tk()
    root.title("Transaction Processor")

    input_dir_var = tk.StringVar(value=DEFAULT_INPUT_DIR)
    output_file_var = tk.StringVar(value=DEFAULT_OUTPUT_FILE)

    style = ttk.Style()
    style.theme_use('clam')  # Modern theme resembling Windows 11

    frame = ttk.Frame(root, padding="20 20 20 20")
    frame.pack(padx=10, pady=10, fill="both", expand=True)

    input_dir_label = ttk.Label(frame, text="Screenshots Directory:")
    input_dir_label.grid(row=0, column=0, pady=5, sticky="W")
    input_dir_entry = ttk.Entry(frame, textvariable=input_dir_var, width=50)
    input_dir_entry.grid(row=0, column=1, pady=5, sticky="EW")
    input_dir_button = ttk.Button(frame, text="Browse", command=select_input_directory)
    input_dir_button.grid(row=0, column=2, pady=5, padx=5)

    output_file_label = ttk.Label(frame, text="Output File:")
    output_file_label.grid(row=1, column=0, pady=5, sticky="W")
    output_file_entry = ttk.Entry(frame, textvariable=output_file_var, width=50)
    output_file_entry.grid(row=1, column=1, pady=5, sticky="EW")
    output_file_button = ttk.Button(frame, text="Browse", command=select_output_file)
    output_file_button.grid(row=1, column=2, pady=5, padx=5)

    process_button = ttk.Button(frame, text="Process Transactions",
                                command=lambda: process_transactions(input_dir_var.get(), output_file_var.get()),
                                width=30)
    process_button.grid(row=2, column=0, columnspan=3, pady=10)

    for child in frame.winfo_children():
        child.grid_configure(padx=5, pady=5)

    sv_ttk.set_theme("dark")
    root.mainloop()


if __name__ == "__main__":
    create_gui()
